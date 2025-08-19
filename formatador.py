from docx import Document
from docx.shared import Pt
from collections import Counter
import re

# Função para apagar um parágrafo do documento
def apagar_paragrafo(paragrafo):
    p = paragrafo._element
    p.getparent().remove(p)
    paragrafo._p = paragrafo._element = None

# Abre o documento original, que será lido e modificado
prova = Document("prova.docx")

apagando = False # Flag que controla quando começar a apagar

# Laço que remove todo o conteúdo a partir da palavra-chave "INÍCIO"
for paragrafo in prova.paragraphs[:]:
    if "INÍCIO" in paragrafo.text:
        apagando = True 
    if apagando:
        apagar_paragrafo(paragrafo)

# Salva o documento sem conteúdo, removido após a palavra-chave "INÍCIO"
prova.save("prova_limpa.docx") 

# Abre o documento já montado com o conteúdo desejado
prova = Document("prova_montada.docx") 

# Laço que itera sobre os parágrafos e padroniza a formatação do texto
for paragrafo in prova.paragraphs:
    paragrafo.paragraph_format.line_spacing = 1.0 
    paragrafo.paragraph_format.space_before = Pt(0) 
    paragrafo.paragraph_format.space_after = Pt(0) 
    for run in paragrafo.runs:
        run.font.name = "Arial" 
        run.font.size = Pt(10) 

# Salva o documento final formatado
prova.save("prova_final.docx") 

# Expressão para identificar alternativas no formato: a) A)
padrao_alternativas = r'^([a-eA-E])\)\s+(.*)' 
alternativas = [] 

# Itera sobre os parágrafos para extrair as alternativas
for paragrafo in prova.paragraphs: 
    texto_paragrafo = paragrafo.text.strip()
    if texto_paragrafo:
        matches = re.findall(padrao_alternativas, texto_paragrafo)
        for match in matches:
            alternativa_texto = match[1].strip() 
            if alternativa_texto:
                alternativas.append(alternativa_texto)

# Conta quantas vezes cada alternativa aparece
contador = Counter(alternativas) 

print('Alternativas: \n')

# Lista todas as alternativas encontradas
for i, alternativa in enumerate(alternativas, 1):
    print(f'{i}. {alternativa}\n')

# Identifica alternativas repetidas
print('Alternativas que se repetem: ')
repetidas = {k: v for k, v in contador.items() if v > 1}

if repetidas:
    for alternativa, qtd in repetidas.items():
        print(f' "{alternativa}" - Repetida {qtd} vez(es)')
else:
    print('Nenhuma alternativa repetida.')